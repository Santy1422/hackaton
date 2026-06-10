"""Endpoint: narrativa de informes redactada por Claude sobre datos reales de la DB.

El backend calcula TODOS los números desde `forecast_13w` + `covenant_rules` y se
los pasa a Claude; el modelo solo escribe la prosa (nunca inventa cifras).
"""

from __future__ import annotations

import io
import secrets
import time
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from db.database import get_connection, query
from models.report_narrative import generate_narrative

from ..auth import get_current_user, require_roles
from ..validation import SCENARIOS, validate_scenario

router = APIRouter(tags=["reports"])

# Store efímero de PDFs generados (1 réplica) → URL pública para que Zavu los baje.
_PDF_STORE: dict[str, tuple[bytes, float]] = {}
_PDF_TTL = 3600  # 1 hora


def publish_pdf(data: bytes) -> str:
    """Guarda un PDF y devuelve un token para servirlo en /api/reports/pdf/{token}."""
    now = time.time()
    for k, (_, exp) in list(_PDF_STORE.items()):
        if exp < now:
            _PDF_STORE.pop(k, None)
    token = secrets.token_urlsafe(16)
    _PDF_STORE[token] = (data, now + _PDF_TTL)
    return token


@router.get("/pdf/{token}")
def get_pdf(token: str):
    """Sirve un PDF generado (público, efímero) — lo descarga Zavu para WhatsApp."""
    item = _PDF_STORE.get(token)
    if not item or item[1] < time.time():
        return Response(status_code=404)
    return Response(
        content=item[0],
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="altis-forecast-{token[:8]}.pdf"'},
    )


@router.get("/download/{scenario}")
def download_pdf(scenario: str, user: dict = Depends(get_current_user)):
    """Genera y devuelve el PDF del informe (server-side, datos reales).

    Robusto: no depende de html2pdf en el browser. Cualquier usuario autenticado.
    """
    validate_scenario(scenario)
    from models.pdf_report import build_pdf

    pdf = build_pdf(scenario)
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="altis-forecast-{scenario}.pdf"'},
    )


SCEN_LABEL = {"base": "Base", "wet_qtr": "Wet quarter", "dry_qtr": "Dry quarter"}


def _eur(n) -> str:
    v = round(float(n or 0))
    return f"-€{abs(v):,.0f}" if v < 0 else f"€{v:,.0f}"


def _scenario_stats(con, scenario: str, threshold: float) -> dict:
    """Agrega forecast_13w por semana y deriva status / headroom / valle (desde la DB)."""
    weeks = query(
        con,
        "SELECT forecast_week, SUM(net_cashflow) AS net, "
        "SUM(d1_milestone_billing) AS d1, SUM(d2_materials_outflow) AS d2, "
        "SUM(d3_subcon_payment) AS d3, SUM(d4_customer_collection) AS d4, "
        "SUM(d5_weather_impact) AS d5 "
        "FROM forecast_13w WHERE scenario = ? GROUP BY forecast_week ORDER BY forecast_week",
        [scenario],
    )
    if not weeks:
        return {}
    cum = 0.0
    min_cum, min_week = None, None
    drivers = {f"d{i}": 0.0 for i in range(1, 6)}
    total_net = 0.0
    for w in weeks:
        net = float(w["net"] or 0)
        total_net += net
        cum += net
        for i in range(1, 6):
            drivers[f"d{i}"] += float(w[f"d{i}"] or 0)
        if min_cum is None or cum < min_cum:
            min_cum, min_week = cum, w["forecast_week"]
    headroom = min_cum - threshold
    status = "BREACH" if min_cum < threshold else "WATCH" if min_cum < threshold + 200000 else "SAFE"
    return {
        "status": status,
        "total_net_cashflow_fmt": _eur(total_net),
        "ending_cash_fmt": _eur(cum),
        "low_point_fmt": _eur(min_cum),
        "low_week": min_week,
        "headroom_fmt": _eur(headroom),
        "drivers_fmt": {k: _eur(v) for k, v in drivers.items()},
    }


class NarrativeBody(BaseModel):
    scenario: str = "base"
    kind: str = "weekly"  # weekly | monthly


@router.post("/narrative")
def report_narrative(
    body: NarrativeBody, user: dict = Depends(require_roles("pe_board", "cfo"))
):
    validate_scenario(body.scenario)
    con = get_connection()
    rules = query(
        con,
        "SELECT value FROM covenant_rules WHERE threshold_type='min_cumulative_cashflow' "
        "ORDER BY id LIMIT 1",
    )
    threshold = float(rules[0]["value"]) if rules else -500000.0

    all_scen = {}
    for s in SCENARIOS:
        st = _scenario_stats(con, s, threshold)
        if st:
            all_scen[s] = {"status": st["status"], "headroom_fmt": st["headroom_fmt"]}
    active = _scenario_stats(con, body.scenario, threshold)
    con.close()

    payload = {
        "kind": body.kind,
        "scenario": body.scenario,
        "scenario_label": SCEN_LABEL.get(body.scenario, body.scenario),
        "covenant_threshold_fmt": _eur(threshold),
        "active": active,
        "all_scenarios": all_scen,
    }

    narrative = generate_narrative(payload)
    return {"scenario": body.scenario, "kind": body.kind, "data": payload, "narrative": narrative}


# ---------------------------------------------------------------------------
# Word document export
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _table_row(table, cells: list[str], bold: bool = False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row


def _build_docx(con, threshold: float) -> bytes:
    doc = Document()

    # ---- Title block -------------------------------------------------------
    title = doc.add_heading("Altis Groep — 13-Week Cash Flow Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Generated {date.today().strftime('%d %B %Y')}  ·  Confidential")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ---- Scenario summary table --------------------------------------------
    _heading(doc, "1. Scenario Summary", 1)
    doc.add_paragraph(
        "Three scenarios are modelled: Base (current run-rate), Wet Quarter "
        "(prolonged rain/frost → delayed billing), and Dry Quarter (accelerated execution)."
    )
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Scenario", "Net CF (13w)", "Ending Cash", "Low Point", "Low Week", "Covenant"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for sc in SCENARIOS:
        st = _scenario_stats(con, sc, threshold)
        if not st:
            continue
        row = tbl.add_row().cells
        row[0].text = SCEN_LABEL.get(sc, sc)
        row[1].text = st["total_net_cashflow_fmt"]
        row[2].text = st["ending_cash_fmt"]
        row[3].text = st["low_point_fmt"]
        row[4].text = f"Week {st['low_week']}"
        row[5].text = st["status"]

    doc.add_paragraph()

    # ---- Per-scenario driver breakdown ------------------------------------
    _heading(doc, "2. Cash Flow Drivers by Scenario", 1)
    DRIVER_LABELS = {
        "d1": "D1 — Milestone billing",
        "d2": "D2 — Materials outflow",
        "d3": "D3 — Subcontractor payments",
        "d4": "D4 — Customer collections",
        "d5": "D5 — Weather impact",
    }
    for sc in SCENARIOS:
        st = _scenario_stats(con, sc, threshold)
        if not st:
            continue
        _heading(doc, f"{SCEN_LABEL.get(sc, sc)} scenario", 2)
        dtbl = doc.add_table(rows=1, cols=2)
        dtbl.style = "Table Grid"
        dtbl.rows[0].cells[0].text = "Driver"
        dtbl.rows[0].cells[1].text = "13-week total"
        for run in dtbl.rows[0].cells[0].paragraphs[0].runs:
            run.bold = True
        for run in dtbl.rows[0].cells[1].paragraphs[0].runs:
            run.bold = True
        for key, label in DRIVER_LABELS.items():
            r = dtbl.add_row().cells
            r[0].text = label
            r[1].text = st["drivers_fmt"].get(key, "—")
        doc.add_paragraph()

    # ---- Weekly forecast detail (base) ------------------------------------
    _heading(doc, "3. Week-by-Week Forecast (Base Scenario)", 1)
    weeks = query(
        con,
        "SELECT forecast_week, week_start, "
        "SUM(net_cashflow) AS net, SUM(gross_inflow) AS inflow, SUM(gross_outflow) AS outflow "
        "FROM forecast_13w WHERE scenario = 'base' "
        "GROUP BY forecast_week, week_start ORDER BY forecast_week",
    )
    if weeks:
        wtbl = doc.add_table(rows=1, cols=5)
        wtbl.style = "Table Grid"
        for i, h in enumerate(["Week", "Start date", "Inflow", "Outflow", "Net CF"]):
            wtbl.rows[0].cells[i].text = h
            for run in wtbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        cum = 0.0
        for w in weeks:
            cum += float(w["net"] or 0)
            r = wtbl.add_row().cells
            r[0].text = str(w["forecast_week"])
            r[1].text = str(w["week_start"])
            r[2].text = _eur(w["inflow"])
            r[3].text = _eur(w["outflow"])
            r[4].text = _eur(w["net"])
    doc.add_paragraph()

    # ---- Covenant status --------------------------------------------------
    _heading(doc, "4. Bank Covenant Status", 1)
    doc.add_paragraph(
        f"Covenant threshold: minimum cumulative cash flow ≥ {_eur(threshold)} "
        "over the 13-week horizon."
    )
    ctbl = doc.add_table(rows=1, cols=3)
    ctbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Worst-case headroom", "Status"]):
        ctbl.rows[0].cells[i].text = h
        for run in ctbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for sc in SCENARIOS:
        st = _scenario_stats(con, sc, threshold)
        if not st:
            continue
        r = ctbl.add_row().cells
        r[0].text = SCEN_LABEL.get(sc, sc)
        r[1].text = st["headroom_fmt"]
        r[2].text = st["status"]
    doc.add_paragraph()

    # ---- Weather & billing insight ----------------------------------------
    _heading(doc, "5. Weather vs Billing — Key Finding", 1)
    doc.add_paragraph(
        "Empirical analysis across 36 months (2023–2025) shows that temperature "
        "explains only 1.4% of billing variance (Pearson r = 0.12, R² = 0.014, p = 0.48). "
        "Weather is NOT the primary driver of revenue movement."
    )
    doc.add_paragraph(
        "The 2024 billing gap traces to large one-off projects (58/59xxx series) completing "
        "without a replacement pipeline, while recurring maintenance contracts (10xxx) "
        "grew throughout the same period."
    )

    # Project categories
    proj_rows = query(
        con,
        "SELECT project_code, year, SUM(credit) AS revenue "
        "FROM transactions "
        "WHERE credit > 0 AND year BETWEEN 2023 AND 2025 "
        "  AND project_code IS NOT NULL "
        "  AND project_code NOT IN ('nan', 'None', '') "
        "GROUP BY project_code, year ORDER BY project_code, year",
    )
    if proj_rows:
        by_type: dict[str, dict[int, float]] = {"recurring": {}, "large_project": {}, "other": {}}
        for r in proj_rows:
            prefix = str(r["project_code"] or "").split(".")[0].strip()
            ptype = "recurring" if prefix.startswith("10") else (
                "large_project" if prefix.startswith(("58", "59")) else "other"
            )
            y = int(r["year"])
            by_type[ptype][y] = by_type[ptype].get(y, 0.0) + float(r["revenue"] or 0)

        ptbl = doc.add_table(rows=1, cols=4)
        ptbl.style = "Table Grid"
        for i, h in enumerate(["Category", "2023", "2024", "2025"]):
            ptbl.rows[0].cells[i].text = h
            for run in ptbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        labels = {"recurring": "Recurring contracts (10xxx)", "large_project": "Large one-off (58/59xxx)"}
        for ptype, label in labels.items():
            rd = ptbl.add_row().cells
            rd[0].text = label
            for i, yr in enumerate([2023, 2024, 2025], 1):
                rd[i].text = _eur(by_type[ptype].get(yr, 0.0))
    doc.add_paragraph()

    # ---- Annual revenue ---------------------------------------------------
    _heading(doc, "6. Historical Revenue by Year", 1)
    annual = query(
        con,
        "SELECT year, opco, SUM(credit) AS revenue "
        "FROM transactions WHERE credit > 0 AND year IS NOT NULL "
        "GROUP BY year, opco ORDER BY year, opco",
    )
    if annual:
        years = sorted({int(r["year"]) for r in annual if r["year"]})
        opcos = sorted({r["opco"] for r in annual})
        idx: dict[tuple, float] = {
            (int(r["year"]), r["opco"]): float(r["revenue"] or 0) for r in annual
        }
        atbl = doc.add_table(rows=1, cols=1 + len(years))
        atbl.style = "Table Grid"
        atbl.rows[0].cells[0].text = "Opco"
        for i, yr in enumerate(years, 1):
            atbl.rows[0].cells[i].text = str(yr)
            for run in atbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for op in opcos:
            rw = atbl.add_row().cells
            rw[0].text = op
            for i, yr in enumerate(years, 1):
                rw[i].text = _eur(idx.get((yr, op), 0.0))
        # totals row
        tot = atbl.add_row().cells
        tot[0].text = "TOTAL"
        for run in tot[0].paragraphs[0].runs:
            run.bold = True
        for i, yr in enumerate(years, 1):
            t = sum(idx.get((yr, op), 0.0) for op in opcos)
            tot[i].text = _eur(t)
            for run in tot[i].paragraphs[0].runs:
                run.bold = True
    doc.add_paragraph()

    # ---- Footer note ------------------------------------------------------
    doc.add_paragraph(
        "This report was generated automatically from the Altis Groep forecast system. "
        "All figures derive from reconciled transaction data (GB Snelstart + Exact). "
        "Scenarios are model outputs — not commitments. For internal use only."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/docx")
def download_report_docx(
    scenario: str = "base",
    user: dict = Depends(require_roles("pe_board", "cfo")),
):
    """Download the full 13-week scenario report as a Word document (.docx)."""
    validate_scenario(scenario)
    con = get_connection()
    try:
        rules = query(
            con,
            "SELECT value FROM covenant_rules WHERE threshold_type='min_cumulative_cashflow' "
            "ORDER BY id LIMIT 1",
        )
        threshold = float(rules[0]["value"]) if rules else -500_000.0
        docx_bytes = _build_docx(con, threshold)
    finally:
        con.close()

    filename = f"altis_13w_report_{date.today().isoformat()}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
